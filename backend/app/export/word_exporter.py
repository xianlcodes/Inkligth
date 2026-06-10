"""
Word (.docx) 导出器
将 Markdown 文本转换为中文学术排版的 Word 文档。

依赖: python-docx (>= 1.2.0)
"""

import re
import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

# ── XML 安全 ──

# XML 1.0 允许的字符范围
_RE_XML_CLEAN = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _clean_xml_text(text: str) -> str:
    """移除 XML 1.0 不允许的字符（PDF 提取常见 \\x0c 换页符等）"""
    if not text:
        return text
    return _RE_XML_CLEAN.sub("", text)

# ── 样式常量 ──

FONT_BODY = "SimSun"        # 正文：宋体
FONT_HEADING = "SimHei"     # 标题：黑体
FONT_CODE = "Consolas"      # 代码：Consolas

SIZE_BODY = Pt(12)          # 正文 12pt
SIZE_H1 = Pt(18)            # 一级标题 18pt
SIZE_H2 = Pt(16)            # 二级标题 16pt
SIZE_H3 = Pt(14)            # 三级标题 14pt
SIZE_CODE = Pt(10)          # 代码 10pt

COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)
COLOR_DARK = RGBColor(0x33, 0x33, 0x33)

LINE_SPACING = 1.5

PAGE_WIDTH = Cm(21.0)       # A4
PAGE_HEIGHT = Cm(29.7)
MARGIN = Inches(1)          # 1 英寸边距


def _set_cell_shading(cell, color: str):
    """设置表格单元格底纹颜色 (e.g. 'D9D9D9')"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_heading_with_font(doc: Document, text: str, level: int):
    """添加带中文字体的标题"""
    heading = doc.add_heading(text, level=level)
    font_size = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3}.get(level, SIZE_BODY)
    for run in heading.runs:
        run.font.name = FONT_HEADING
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        run.font.size = font_size
        run.font.color.rgb = COLOR_DARK
        run.font.bold = True
    return heading


def _add_paragraph_with_font(doc: Document, text: str, style=None):
    """添加正文段落（宋体 12pt，首行缩进 2 字符）"""
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.first_line_indent = Pt(24)  # ~2 字符
    p.paragraph_format.line_spacing = LINE_SPACING
    for run in p.runs:
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        run.font.size = SIZE_BODY
    return p


def _add_code_block(doc: Document, code: str):
    """添加代码块（Consolas 10pt，灰色背景）"""
    lines = code.strip().split("\n")
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        # 灰色背景（仅首行和末行设背景，减少 xml 膨胀）
        run = p.add_run(line if line else " ")
        run.font.name = FONT_CODE
        run.font.size = SIZE_CODE
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # 添加灰色底纹
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        run._element.get_or_add_rPr().append(shading)


def _add_blockquote(doc: Document, text: str):
    """添加引用块（灰色文字，左侧竖线边框）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = LINE_SPACING
    # 左侧竖线（通过段落边框实现）
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:left w:val="single" w:sz="12" w:space="8" w:color="999999"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run.font.size = SIZE_BODY
    run.font.color.rgb = COLOR_GRAY
    run.font.italic = True


def _add_table_from_markdown(doc: Document, rows: list[list[str]], header: bool = True):
    """从 Markdown 表格数据创建 Word 表格"""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            run.font.name = FONT_BODY
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
            run.font.size = Pt(10)
            if header and i == 0:
                run.font.bold = True
                _set_cell_shading(cell, "D9D9D9")


def _parse_inline_formatting(text: str) -> list[tuple[str, bool, bool, bool]]:
    """
    解析内联 Markdown 格式，返回 [(文本, 粗体, 斜体, 代码), ...]
    支持: **bold**, *italic*, `code`, ~~strikethrough~~
    """
    segments = []
    pattern = r"(\*\*(.+?)\*\*)|(\*(.+?)\*)|(`(.+?)`)|(~~(.+?)~~)"
    last_end = 0
    for m in re.finditer(pattern, text):
        if m.start() > last_end:
            segments.append((text[last_end:m.start()], False, False, False))
        if m.group(2):
            segments.append((m.group(2), True, False, False))
        elif m.group(4):
            segments.append((m.group(4), False, True, False))
        elif m.group(6):
            segments.append((m.group(6), False, False, True))
        elif m.group(8):
            segments.append((m.group(8), False, False, False))
        last_end = m.end()
    if last_end < len(text):
        segments.append((text[last_end:], False, False, False))
    return segments


def _render_inline_run(paragraph, text: str, bold: bool = False, italic: bool = False, is_code: bool = False):
    """将格式化文本渲染到段落"""
    run = paragraph.add_run(text)
    if is_code:
        run.font.name = FONT_CODE
        run.font.size = SIZE_CODE
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    else:
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        run.font.size = SIZE_BODY
    run.font.bold = bold
    run.font.italic = italic
    return run


def markdown_to_docx(
    markdown_text: str,
    output_path: str,
    *,
    title: str = "",
    include_toc: bool = False,
    page_numbers: bool = True,
) -> str:
    """
    将 Markdown 文本转换为 Word 文档。

    Args:
        markdown_text: Markdown 格式的源文本
        output_path: 输出 .docx 文件路径
        title: 文档标题
        include_toc: 是否插入目录
        page_numbers: 是否显示页码

    Returns:
        输出文件的路径
    """
    doc = Document()

    # 清理 XML 非法字符（PDF 提取常含 \\x0c 等）
    markdown_text = _clean_xml_text(markdown_text)
    title = _clean_xml_text(title)

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # ── 标题 ──
    if title:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(title)
        title_run.font.name = FONT_HEADING
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = COLOR_DARK
        title_p.paragraph_format.space_after = Pt(24)

    # ── 目录 ──
    if include_toc:
        toc_p = doc.add_paragraph()
        toc_run = toc_p.add_run("目 录")
        toc_run.font.name = FONT_HEADING
        toc_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        toc_run.font.size = SIZE_H2
        toc_run.font.bold = True
        # 添加 TOC 占位符（Word 打开后会提示更新域）
        # 插入一个简单的列表占位，用户可按 Ctrl+A → F9 刷新
        toc_field = parse_xml(
            f'<w:p {nsdecls("w")}>'
            '  <w:fldSimple w:instr=" TOC \\o \\h \\z ">'
            '    <w:r><w:t>[目录 - 请在 Word 中按 Ctrl+A → F9 刷新]</w:t></w:r>'
            '  </w:fldSimple>'
            '</w:p>'
        )
        doc._element.body.append(toc_field)
        doc.add_page_break()

    # ── 解析 Markdown ──
    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            if in_table and table_data:
                _add_table_from_markdown(doc, table_data, header=True)
                table_data = []
                in_table = False
            i += 1
            continue

        # 标题
        h_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2)
            p = _add_heading_with_font(doc, text, level)
            i += 1
            continue

        # 代码块（``` 包裹）
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, "\n".join(code_lines))
            i += 1  # skip closing ```
            continue

        # 水平分割线
        if re.match(r"^[-*_]{3,}\s*$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 50)
            run.font.color.rgb = COLOR_GRAY
            run.font.size = Pt(8)
            i += 1
            continue

        # 引用块
        if stripped.startswith("> "):
            quote_text = stripped[2:]
            _add_blockquote(doc, quote_text)
            i += 1
            continue

        # Markdown 表格（检测 | 分隔行）
        if "|" in stripped and re.match(r"^\|.*\|$", stripped):
            # 跳过对齐行（|---|）
            if re.match(r"^\|[\s:-]+\|$", stripped):
                i += 1
                continue
            if not in_table:
                table_data = []
                in_table = True
            table_data.append([c.strip() for c in stripped.strip("|").split("|")])
            i += 1
            continue

        # 列表项（- 或 * 或 1.）
        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)$", stripped)
        if list_match:
            indent = list_match.group(1)
            text = list_match.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1 + len(indent) * 0.5)
            p.paragraph_format.line_spacing = LINE_SPACING
            for seg_text, bold, italic, is_code in _parse_inline_formatting(text):
                _render_inline_run(p, seg_text, bold, italic, is_code)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = LINE_SPACING
        for seg_text, bold, italic, is_code in _parse_inline_formatting(stripped):
            _render_inline_run(p, seg_text, bold, italic, is_code)
        i += 1

    # 处理未闭合的表格
    if in_table and table_data:
        _add_table_from_markdown(doc, table_data, header=True)

    # ── 页码 ──
    if page_numbers:
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 插入页码域
            run = p.add_run()
            fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._element.append(fld_char_begin)
            run2 = p.add_run()
            instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run2._element.append(instr)
            run3 = p.add_run()
            fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run3._element.append(fld_char_end)

    # ── 保存 ──
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path_obj))
    logger.info("Word document saved: %s", output_path)
    return output_path
